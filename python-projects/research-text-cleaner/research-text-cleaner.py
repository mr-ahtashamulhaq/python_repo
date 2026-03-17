import re
import sys
from pathlib import Path

# Regexes
BRACKET_RE = re.compile(r"\[[^\]]*\]")  # removes anything inside [...] including commas, ranges, etc.
STARS_HASH_RE = re.compile(r"[*#]")     # removes '*' and '#'

def clean_line_preserve(content_line: str) -> str:
    """Clean a single line (no newline char). Returns cleaned line (no trailing newline)."""
    # Remove bracketed citations
    s = BRACKET_RE.sub("", content_line)
    # Remove '*' and '#'
    s = STARS_HASH_RE.sub("", s)
    # Collapse multiple spaces/tabs into a single space (keeps words intact)
    s = re.sub(r"[ \t]+", " ", s)
    # Trim leading/trailing spaces for neatness but keep empty lines empty
    return s.strip() if s.strip() != "" else ""

def clean_text_preserve_newlines(text: str) -> str:
    """
    Clean full text while preserving newline structure.
    Works line-by-line so newline chars remain as in the original.
    """
    lines = text.splitlines(keepends=True)
    out_lines = []
    for line in lines:
        # detect newline part and content part
        if line.endswith("\r\n"):
            newline = "\r\n"
            content = line[:-2]
        elif line.endswith("\n"):
            newline = "\n"
            content = line[:-1]
        else:
            newline = ""
            content = line
        cleaned = clean_line_preserve(content)
        out_lines.append(cleaned + newline)
    return "".join(out_lines)

def process_txt(in_path: Path, out_path: Path, encoding="utf-8"):
    with in_path.open("r", encoding=encoding) as f:
        raw = f.read()
    cleaned = clean_text_preserve_newlines(raw)
    with out_path.open("w", encoding=encoding) as f:
        f.write(cleaned)

def process_docx(in_path: Path, out_path: Path):
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx is required to process .docx files: pip install python-docx")
    doc = Document(in_path)

    # Clean paragraphs
    for p in doc.paragraphs:
        cleaned = clean_text_preserve_newlines(p.text)
        # Note: setting p.text will replace runs/formatting (ok for cleanup)
        p.text = cleaned

    # Clean table cell contents (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Rebuild the cell text from cleaned paragraphs to preserve line breaks inside a cell
                cleaned_pars = [clean_text_preserve_newlines(par.text).rstrip("\n") for par in cell.paragraphs]
                cell.text = "\n".join(cleaned_pars)

    doc.save(out_path)

def main():
    import argparse
    parser = argparse.ArgumentParser(description="Clean research text: remove '*', '#', and [citations]. Preserves newlines.")
    parser.add_argument("input", help="input file (.txt or .docx)")
    parser.add_argument("output", help="output file path")
    args = parser.parse_args()

    inp = Path(args.input)
    out = Path(args.output)

    if not inp.exists():
        print(f"ERROR: input file not found: {inp}", file=sys.stderr)
        sys.exit(2)

    suffix = inp.suffix.lower()
    if suffix == ".docx":
        process_docx(inp, out)
    else:
        process_txt(inp, out)

    print(f"Saved cleaned file to: {out}")

if __name__ == "__main__":
    main()
